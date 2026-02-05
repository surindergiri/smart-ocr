from docx import Document
import os

def create_word_doc(text_pages: list[list[str]], output_path: str):
    """
    Creates a .docx file from the extracted text.
    text_pages: list of pages, where each page is a list of text lines.
    """
    doc = Document()
    doc.add_heading('OCR Extraction Result', 0)

    for i, page_lines in enumerate(text_pages):
        if i > 0:
            doc.add_page_break()
        
        doc.add_heading(f'Page {i+1}', level=1)
        
        # Simple paragraph addition. 
        # For better formatting, we might want to group lines or preserve layout,
        # but for this MVP, we just dump the lines.
        for line in page_lines:
            doc.add_paragraph(line)
            
    doc.save(output_path)
    return output_path
